from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("entregaveis")
OUTPUT_FILE = OUTPUT_DIR / "Proposta_Comercial_HF_Veiculos.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE_GRAY = "F4F6F9"
LIGHT_GRAY = "F2F4F7"
BORDER = "D7DEE8"
BLACK = "000000"
CONTENT_WIDTH = 9360


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(paragraph, before=0, after=8, line=1.333, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = alignment


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:{}".format(key)), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:{}".format(margin)))
        if node is None:
            node = OxmlElement("w:{}".format(margin))
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr

    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    grid_cols = list(grid)
    for index, width in enumerate(widths):
        if index < len(grid_cols):
            grid_cols[index].set(qn("w:w"), str(width))

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=MUTED)


def add_text(paragraph, text, size=11, color=BLACK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    add_text(p, text)
    return p


def add_label_body(doc, label, text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    add_text(p, label, bold=True, color=NAVY)
    add_text(p, text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, text)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, text)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE_GRAY)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "4", "color": BORDER},
        bottom={"val": "single", "sz": "4", "color": BORDER},
        left={"val": "single", "sz": "4", "color": BORDER},
        right={"val": "single", "sz": "4", "color": BORDER},
    )
    p = cell.paragraphs[0]
    set_paragraph_format(p, before=0, after=0, line=1.2, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_text(p, title + " ", size=10.5, color=NAVY, bold=True)
    add_text(p, text, size=10.5, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)


def add_section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    add_text(p, text, size=16, color=BLUE, bold=True)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    add_text(p, text, size=13, color=BLUE, bold=True)
    return p


def write_cell(cell, text, *, size=9, bold=False, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_format(p, before=0, after=0, line=1.12, alignment=alignment)
    add_text(p, text, size=size, color=color, bold=bold)


def add_comparison_table(doc):
    rows = [
        ("Solução", "Formulário digital", "App móvel com AppSheet", "Sistema próprio / PWA"),
        ("Ideal para", "Validar a ideia com rapidez", "Operação recorrente com melhor custo-benefício", "Maior controle e evolução sob medida"),
        ("Recursos-chave", "Checklist, fotos, planilha, Drive e PDF simples", "Cadastro, entrega/devolução, fotos, assinatura, PDF e histórico", "Login, painel, banco de dados, backup e evolução futura"),
        ("Prazo estimado", "5 a 10 dias úteis", "10 a 20 dias úteis", "30 a 60 dias úteis"),
        ("Implantação", "R$ 800 a R$ 1.500", "R$ 2.000 a R$ 4.500", "R$ 5.000 a R$ 12.000+"),
        ("Custo recorrente", "Domínio e suporte opcional", "Domínio, possível licença AppSheet por usuário/mês e suporte", "Domínio, infraestrutura gerenciada e suporte"),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1660, 2530, 2660, 2510])
    headers = ("Critério", "Opção básica", "Opção recomendada", "Opção profissional")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE_GRAY)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            left={"val": "single", "sz": "6", "color": BORDER},
            right={"val": "single", "sz": "6", "color": BORDER},
        )
        write_cell(cell, header, size=8.5, bold=True, color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                left={"val": "single", "sz": "4", "color": BORDER},
                right={"val": "single", "sz": "4", "color": BORDER},
            )
            if index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            write_cell(cell, value, size=8.35, bold=(index == 0), color=NAVY if index == 0 else BLACK,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_investment_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [1880, 2040, 3020, 2420])
    headers = ("Opção", "Implantação estimada", "Custos recorrentes", "Suporte mensal opcional")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE_GRAY)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            left={"val": "single", "sz": "6", "color": BORDER},
            right={"val": "single", "sz": "6", "color": BORDER},
        )
        write_cell(cell, header, size=8.4, bold=True, color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])

    data = [
        ("Básica", "R$ 800 a R$ 1.500", "Domínio: cerca de R$ 40/ano. Armazenamento extra, se necessário, é contratado pelo cliente.", "R$ 100/mês"),
        ("Recomendada", "R$ 2.000 a R$ 4.500", "Domínio + possível licença AppSheet por usuário/mês, conforme plano e quantidade de pessoas.", "R$ 250/mês"),
        ("Profissional", "R$ 5.000 a R$ 12.000+", "Domínio + hospedagem, banco de dados e serviços necessários, definidos no escopo.", "A partir de R$ 400/mês"),
    ]
    for values in data:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                left={"val": "single", "sz": "4", "color": BORDER},
                right={"val": "single", "sz": "4", "color": BORDER},
            )
            if index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            write_cell(cell, value, size=8.3, bold=(index == 0), color=NAVY if index == 0 else BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_delivery_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2500, 2380, 4480])
    headers = ("Opção", "Prazo estimado", "Marco de entrega")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE_GRAY)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            left={"val": "single", "sz": "6", "color": BORDER},
            right={"val": "single", "sz": "6", "color": BORDER},
        )
        write_cell(cell, header, size=8.5, bold=True, color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])

    data = [
        ("Básica", "5 a 10 dias úteis", "Formulário validado, organização no Drive/Sheets e PDF simples."),
        ("Recomendada", "10 a 20 dias úteis", "App de vistoria testado com fluxo de entrega e devolução."),
        ("Profissional", "30 a 60 dias úteis", "PWA publicada, painel, cadastros, banco de dados e backup inicial."),
    ]
    for values in data:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                left={"val": "single", "sz": "4", "color": BORDER},
                right={"val": "single", "sz": "4", "color": BORDER},
            )
            if index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            write_cell(cell, value, size=8.5, bold=(index == 0), color=NAVY if index == 0 else BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)

    if "Small Note" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    add_text(header_p, "PROPOSTA COMERCIAL | VISTORIA VEICULAR", size=8.5, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_after = Pt(0)
    add_text(footer_p, "Winglisson M. Gonçalves | Digitalização de vistoria veicular | Página ", size=8.5, color=MUTED)
    add_page_field(footer_p)

    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.company = ""
    props.title = "Proposta Comercial - Digitalização de Vistoria Veicular"
    props.subject = "Digitalização do processo de vistoria"
    props.comments = ""


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_text(p, "PROPOSTA COMERCIAL", size=12, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    add_text(p, "Digitalização do Processo", size=26, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    add_text(p, "de Vistoria Veicular", size=26, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(20)
    add_text(
        p,
        "Uma solução móvel para registrar entrega e devolução de veículos com checklist, fotos, aceite e PDF.",
        size=12,
        color=MUTED,
    )

    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, [4680, 4680])
    metadata = (
        ("CLIENTE", "HF Veículos"),
        ("PROPONENTE", "Winglisson M. Gonçalves"),
        ("DATA", "[Data de emissão]"),
        ("VALIDADE", "15 dias"),
    )
    for index, (label, value) in enumerate(metadata):
        cell = table.cell(index // 2, index % 2)
        set_cell_shading(cell, LIGHT_BLUE_GRAY)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "4", "color": BORDER},
            bottom={"val": "single", "sz": "4", "color": BORDER},
            left={"val": "single", "sz": "4", "color": BORDER},
            right={"val": "single", "sz": "4", "color": BORDER},
        )
        p = cell.paragraphs[0]
        set_paragraph_format(p, before=0, after=1, line=1.1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(p, label, size=8, color=MUTED, bold=True)
        p = cell.add_paragraph()
        set_paragraph_format(p, before=0, after=0, line=1.1, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_text(p, value, size=10.5, color=NAVY, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build_document():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_section_heading(doc, "1. Contexto e objetivo")
    add_body(
        doc,
        "Hoje, o Termo de Vistoria é preenchido manualmente em PDF. Esse processo toma tempo da equipe, dificulta a leitura e o arquivo das informações e torna mais trabalhoso localizar evidências de uma entrega ou devolução. A proposta é levar a vistoria para o celular, com um fluxo padronizado, fotos organizadas e PDF final disponível para consulta."
    )
    add_body(
        doc,
        "A solução registrará os dados da locadora e do locatário, identificação do veículo, quilometragem, combustível, acessórios, condições externas e internas, pneus, observações, declaração e aceite. O resultado esperado é uma operação mais rápida, rastreável e simples de administrar."
    )

    add_section_heading(doc, "2. Benefícios para a operação")
    benefits = [
        "Vistorias preenchidas no celular, sem depender de papel ou impressão.",
        "Checklists padronizados, reduzindo esquecimentos e divergências no atendimento.",
        "Fotos anexadas à vistoria para registrar as condições do veículo.",
        "PDF gerado para compartilhamento e arquivamento após o atendimento.",
        "Histórico mais fácil de localizar por placa, cliente ou data, conforme a opção escolhida.",
        "Dados organizados para apoiar gestão, conferência e atendimento ao cliente.",
    ]
    for benefit in benefits:
        add_bullet(doc, benefit)

    doc.add_page_break()

    add_section_heading(doc, "3. Comparativo das opções")
    add_body(
        doc,
        "As três alternativas abaixo atendem ao objetivo de digitalizar a vistoria. A escolha depende do estágio atual da operação, da necessidade de histórico e do nível de controle desejado."
    )
    add_comparison_table(doc)

    add_section_heading(doc, "4. Detalhamento das opções")
    add_subheading(doc, "Opção básica - Formulário digital")
    add_body(doc, "Indicada para colocar o processo em prática rapidamente e validar a rotina digital com menor investimento inicial.")
    basic_items = [
        "Sugestão e configuração inicial de domínio próprio .com.br.",
        "Formulário de vistoria responsivo, acessível pelo celular.",
        "Campos para dados da locadora, locatário, veículo, quilometragem, combustível, observações e declaração.",
        "Checklist de acessórios, condições externas, internas e pneus.",
        "Upload de fotos e organização dos registros no Google Drive.",
        "Respostas centralizadas em planilha do Google Sheets.",
        "PDF simples da vistoria para arquivamento e envio.",
    ]
    for item in basic_items:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Limite da opção básica:",
        "é uma solução baseada em formulário. Não inclui painel administrativo, histórico avançado por placa/cliente, assinatura capturada na tela ou aplicativo próprio."
    )

    doc.add_page_break()

    add_subheading(doc, "Opção recomendada - App de vistoria com AppSheet")
    add_body(
        doc,
        "É a alternativa indicada para uma operação de locação que precisa de fluxo completo, histórico organizado e uso recorrente pela equipe, com investimento menor e entrega mais rápida que um sistema totalmente sob medida."
    )
    recommended_items = [
        "Domínio próprio com direcionamento profissional de acesso.",
        "Aplicativo de vistoria acessível pelo celular.",
        "Cadastro de veículos e consulta por placa.",
        "Registro de vistoria de entrega e de devolução.",
        "Checklist completo de acessórios, condições externas, internas e pneus.",
        "Fotos obrigatórias em pontos definidos do processo.",
        "Registro de quilometragem, combustível, observações e declaração.",
        "Assinatura simples coletada na tela como aceite operacional.",
        "Geração de PDF e armazenamento organizado no Google Drive.",
        "Histórico pesquisável por placa e cliente em Google Sheets/Drive.",
    ]
    for item in recommended_items:
        add_bullet(doc, item)
    add_callout(
        doc,
        "Licenciamento AppSheet:",
        "pode haver custo mensal por usuário, de acordo com o plano escolhido e a quantidade de pessoas que utilizarão o app. Esse valor será confirmado após mapear os usuários e poderá ser contratado diretamente pelo cliente."
    )

    add_subheading(doc, "Opção profissional - Sistema próprio / PWA")
    add_body(
        doc,
        "Indicada quando a locadora deseja uma solução exclusiva, com mais autonomia, gestão de usuários e uma base preparada para evoluir junto com a operação.")
    professional_items = [
        "Domínio próprio e sistema web instalável no celular (PWA).",
        "Login de usuários e controle inicial de acessos.",
        "Painel administrativo para veículos, clientes e vistorias.",
        "Cadastro e consulta de veículos, clientes, entregas e devoluções.",
        "Upload de fotos, assinatura simples na tela e geração automática de PDF.",
        "Histórico completo das vistorias em banco de dados.",
        "Rotina inicial de backup e organização de dados.",
        "Base preparada para futuras evoluções, como relatórios, integrações e alertas, mediante novo escopo.",
    ]
    for item in professional_items:
        add_bullet(doc, item)

    doc.add_page_break()

    add_section_heading(doc, "5. Investimento estimado")
    add_body(
        doc,
        "Os valores abaixo são faixas de referência e serão ajustados após a definição de usuários, quantidade de campos, layout do PDF, volume de dados e eventuais integrações. A contratação do domínio .com.br custa aproximadamente R$ 40 por ano."
    )
    add_investment_table(doc)
    add_callout(
        doc,
        "Importante:",
        "o custo de implantação é separado dos custos recorrentes de plataformas, licenças, hospedagem ou armazenamento que forem necessários em cada opção."
    )

    add_section_heading(doc, "6. O que está incluso")
    included = [
        "Levantamento inicial do Termo de Vistoria atual e organização dos campos principais.",
        "Sugestão de domínio e orientação para registro preferencialmente em nome/CNPJ do cliente.",
        "Configuração inicial do ambiente correspondente à opção escolhida.",
        "Construção e testes do fluxo descrito nesta proposta.",
        "Uma sessão remota de apresentação e orientação de uso para a equipe.",
        "Até 2 rodadas de ajustes sobre o escopo aprovado. Cada rodada consolida as solicitações enviadas pelo cliente.",
        "Suporte inicial de 30 dias após a entrega para correção de defeitos diretamente ligados ao escopo contratado.",
    ]
    for item in included:
        add_bullet(doc, item)

    add_section_heading(doc, "7. O que não está incluso")
    excluded = [
        "Registro ou renovação anual de domínio, licenças AppSheet, Google Workspace, armazenamento, hospedagem, banco de dados ou outros serviços de terceiros.",
        "Assinatura digital avançada, certificado ICP-Brasil ou garantia de validade jurídica específica.",
        "Integrações com sistemas de locação, ERP, WhatsApp, pagamentos, SMS ou APIs externas, salvo orçamento complementar.",
        "Migração em massa de vistorias antigas, digitalização de arquivos físicos ou tratamento de base histórica.",
        "Criação de identidade visual, logotipo, fotografia profissional ou produção de conteúdo institucional.",
        "Novas funcionalidades, alterações estruturais ou ajustes solicitados após as 2 rodadas incluídas.",
        "Suporte recorrente após o período inicial, salvo contratação do plano mensal opcional.",
    ]
    for item in excluded:
        add_bullet(doc, item)

    doc.add_page_break()

    add_section_heading(doc, "8. Prazo estimado de entrega")
    add_delivery_table(doc)
    add_body(
        doc,
        "Os prazos começam após o recebimento do termo atual, informações necessárias, definição do domínio, acessos às contas envolvidas e aprovação do fluxo. Atrasos no envio de materiais, acessos ou aprovações deslocam o cronograma na mesma proporção."
    )

    add_section_heading(doc, "9. Condições e observações importantes")
    observations = [
        "Condição de pagamento sugerida: 50% na aprovação/início e 50% na disponibilização para validação final. A condição pode ser ajustada antes da contratação.",
        "A proposta é válida por 15 dias. O início ocorre após confirmação da opção, aprovação comercial e recebimento dos itens necessários.",
        "O domínio e as contas principais devem ficar, preferencialmente, em nome ou CNPJ do cliente. Será necessário conceder acesso técnico para a configuração e manutenção.",
        "AWS não é necessária no início. Nas opções básica e recomendada, a solução pode operar com recursos Google e AppSheet. Na opção profissional, a infraestrutura será escolhida conforme a necessidade real da operação.",
        "A assinatura registrada na tela é um aceite simples para controle operacional. Não substitui, por si só, uma assinatura digital avançada nem é oferecida com promessa de validade jurídica específica.",
        "Qualquer evolução fora do escopo aprovado será analisada e orçada antes da execução.",
    ]
    for item in observations:
        add_bullet(doc, item)

    add_section_heading(doc, "10. Recomendação final")
    add_callout(
        doc,
        "Recomendação: Opção AppSheet.",
        "Ela oferece o melhor custo-benefício para o cenário atual: permite cadastro, entrega e devolução, fotos obrigatórias, assinatura simples, PDF e histórico, sem exigir o investimento inicial de um sistema próprio. É uma forma segura de estruturar a operação e aprender com o uso real antes de uma evolução maior."
    )

    add_section_heading(doc, "11. Próximos passos para iniciar")
    next_steps = [
        "Escolher a opção que melhor atende à operação atual.",
        "Enviar o Termo de Vistoria atual e indicar os responsáveis pelo uso e aprovação.",
        "Definir o domínio desejado e autorizar o registro/configuração em nome do cliente.",
        "No caso do AppSheet, confirmar a quantidade de usuários e os e-mails que terão acesso.",
        "Aprovar o escopo comercial para iniciar o mapeamento do fluxo e o cronograma.",
    ]
    for item in next_steps:
        add_numbered(doc, item)

    add_section_heading(doc, "Versão curta para WhatsApp")
    add_callout(
        doc,
        "Mensagem sugerida:",
        "Olá, equipe da HF Veículos. Preparei uma proposta para digitalizar as vistorias pelo celular, com checklist, fotos, aceite simples, PDF e histórico. A opção que recomendo é o AppSheet: implantação estimada entre R$ 2.000 e R$ 4.500, mais possível licença por usuário/mês e domínio em torno de R$ 40/ano. O prazo é de 10 a 20 dias úteis após os materiais e acessos. Se preferirem, também há uma opção básica (R$ 800 a R$ 1.500) e um sistema próprio (a partir de R$ 5.000). Com a escolha, seguimos para o mapeamento e início da implantação."
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE.resolve())


if __name__ == "__main__":
    build_document()
